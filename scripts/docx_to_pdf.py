"""
docx_to_pdf.py — Conversion DOCX → PDF via Microsoft Word (COM automation)
1. Supprime les paragraphes vides en fin de section (pages blanches)
2. Enregistre en PDF dans outputs/Cyclistic_Report.pdf
"""

import os
import time
import comtypes.client
from docx import Document
from docx.oxml.ns import qn

BASE_DIR  = r"C:\Users\Republic Of Computer\Desktop\Cyclistic_Project"
DOCX_IN   = os.path.join(BASE_DIR, "06_report.docx")
PDF_OUT   = os.path.join(BASE_DIR, "06_report.pdf")
DOCX_CLEAN = os.path.join(BASE_DIR, "06_report_clean.docx")

# ── Étape 1 : Supprimer les pages blanches du DOCX ───────────────────────────

def is_empty_paragraph(para):
    """Retourne True si le paragraphe est vide (aucun texte, aucune image)."""
    text = para.text.strip()
    if text:
        return False
    # Vérifier s'il contient des images ou autres éléments
    if para._element.findall('.//' + qn('a:blip')):
        return False
    if para._element.findall('.//' + qn('wp:inline')):
        return False
    return True

def has_page_break(para):
    """Retourne True si le paragraphe contient un saut de page."""
    for run in para.runs:
        for br in run._element.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
    # Vérifier dans les propriétés de paragraphe
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pageBreak = pPr.find(qn('w:pageBreakBefore'))
        if pageBreak is not None and pageBreak.get(qn('w:val'), 'true') != 'false':
            return True
    return False

print("Étape 1 — Chargement du DOCX...")
doc = Document(DOCX_IN)

# Supprimer les paragraphes vides consécutifs (max 1 vide autorisé entre sections)
paras = doc.paragraphs
removed = 0
i = 0
while i < len(paras) - 1:
    para = paras[i]
    next_para = paras[i + 1]
    # Si deux paragraphes vides consécutifs → supprimer le premier
    if is_empty_paragraph(para) and is_empty_paragraph(next_para):
        p = para._element
        p.getparent().remove(p)
        paras = doc.paragraphs  # Rafraîchir la liste
        removed += 1
    else:
        i += 1

print(f"  {removed} paragraphes vides supprimés.")

# Sauvegarder le DOCX nettoyé
doc.save(DOCX_CLEAN)
print(f"  DOCX nettoyé sauvegardé : {DOCX_CLEAN}")

# ── Étape 2 : Conversion en PDF via Word COM ──────────────────────────────────

print("\nÉtape 2 — Conversion en PDF via Microsoft Word...")

wdFormatPDF = 17  # Constante Word pour export PDF

word = None
doc_word = None
try:
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    word.DisplayAlerts = False

    doc_word = word.Documents.Open(DOCX_CLEAN)
    time.sleep(1)  # Laisser Word charger complètement

    doc_word.SaveAs(PDF_OUT, FileFormat=wdFormatPDF)
    doc_word.Close(False)
    print(f"  PDF généré : {PDF_OUT}")

except Exception as e:
    print(f"  ERREUR : {e}")
    raise
finally:
    if word:
        word.Quit()

# ── Étape 3 : Nettoyage et vérification ──────────────────────────────────────

# Supprimer le DOCX temporaire nettoyé
if os.path.exists(DOCX_CLEAN):
    os.remove(DOCX_CLEAN)

size_kb = os.path.getsize(PDF_OUT) // 1024
print(f"\nOK — PDF final : {PDF_OUT}")
print(f"     Taille    : {size_kb} KB")
print(f"     Encodage  : natif Word (Unicode complet)")
