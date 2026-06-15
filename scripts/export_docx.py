"""
export_docx.py — Rapport professionnel complet Cyclistic (.docx) v2
Auteur : Sidi Mohamed ALLY — Analyste de données Professionnel Certifié
"""

import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

BASE_DIR  = r"C:\Users\Republic Of Computer\Desktop\Cyclistic_Project"
FIG_DIR   = os.path.join(BASE_DIR, "figures")
SQL_FILE  = os.path.join(BASE_DIR, "scripts", "03_queries.sql")
PY_CLEAN  = os.path.join(BASE_DIR, "scripts", "02_cleaning.py")
PY_ANAL   = os.path.join(BASE_DIR, "scripts", "03_analysis.py")
OUT_DOCX  = os.path.join(BASE_DIR, "06_report.docx")
BADGE     = os.path.join(FIG_DIR, "badge_google_da.png")
CREDLY    = "https://www.credly.com/go/RarXbDJG"

# ── Palette ────────────────────────────────────────────────────────────────────
BLUE  = RGBColor(0x1A, 0x73, 0xE8)
GREEN = RGBColor(0x34, 0xA8, 0x53)
AMBER = RGBColor(0xB0, 0x60, 0x00)
DARK  = RGBColor(0x20, 0x21, 0x24)
GRAY  = RGBColor(0x5F, 0x63, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xF1, 0xF3, 0xF4)

# ── Helpers XML ────────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#').upper())
    tcPr.append(shd)

def bottom_border(para, color='1A73E8'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_page_num(para):
    run = para.add_run()
    for ftype in ['begin', None, 'end']:
        if ftype in ['begin', 'end']:
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), ftype)
            run._r.append(fc)
        else:
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = ' PAGE '
            run._r.append(instr)
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

def add_hyperlink(para, text, url, size=9, color='1A73E8'):
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u');   u.set(qn('w:val'), 'single')
    clr = OxmlElement('w:color'); clr.set(qn('w:val'), color)
    sz  = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size * 2))
    rPr.extend([u, clr, sz])
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)

# ── Builders ───────────────────────────────────────────────────────────────────
def body(doc, text, size=10, space_after=6, color=None, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color or DARK
    r.font.bold = bold
    r.font.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)

def h1(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{num}  ")
    r1.font.size = Pt(16); r1.font.bold = True; r1.font.color.rgb = BLUE
    r2 = p.add_run(title)
    r2.font.size = Pt(16); r2.font.bold = True; r2.font.color.rgb = DARK
    bottom_border(p)

def h2(doc, label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"{label}  {title}")
    r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = DARK

def figure(doc, fname, fig_num, caption_text, width=14.0):
    path = os.path.join(FIG_DIR, fname)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f"Fig. {fig_num} — {caption_text}  ({fname})")
    r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = GRAY
    cp.paragraph_format.space_after = Pt(10)

def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        set_bg(cell, '#1A73E8')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            if ri % 2 == 0: set_bg(cell, '#F8F9FA')
            r = cell.paragraphs[0].add_run(str(val)); r.font.size = Pt(9)
    if widths:
        for row_obj in tbl.rows:
            for ci, cell in enumerate(row_obj.cells):
                if ci < len(widths): cell.width = Cm(widths[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def insight_box(doc, text):
    """Encadré 'Insight clé' gris clair en fin de sous-section."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_bg(cell, '#E8F0FE')
    ph = cell.paragraphs[0]
    rh = ph.add_run("  Insight clé")
    rh.font.bold = True; rh.font.size = Pt(9); rh.font.color.rgb = BLUE
    pb = cell.add_paragraph()
    rb = pb.add_run(f"  {text}")
    rb.font.size = Pt(10); rb.font.color.rgb = DARK; rb.font.italic = True
    pb.paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def colored_box(doc, header, bg_hex, items, hdr_color):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_bg(cell, bg_hex)
    ph = cell.paragraphs[0]
    rh = ph.add_run(f"  {header}")
    rh.font.size = Pt(11); rh.font.bold = True; rh.font.color.rgb = hdr_color
    for lbl, val in items:
        pi = cell.add_paragraph()
        pi.paragraph_format.left_indent = Cm(0.4)
        pi.paragraph_format.space_after = Pt(3)
        if lbl:
            rl = pi.add_run(f"{lbl} : "); rl.font.bold = True
            rl.font.size = Pt(10); rl.font.color.rgb = DARK
        rv = pi.add_run(val); rv.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def code_block(doc, text, title=None):
    if title:
        pt = doc.add_paragraph()
        pt.paragraph_format.space_before = Pt(8)
        rt = pt.add_run(title)
        rt.font.bold = True; rt.font.size = Pt(9); rt.font.color.rgb = BLUE
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(7.5); r.font.color.rgb = DARK

# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Marges
sec0 = doc.sections[0]
for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
    setattr(sec0, attr, Cm(2.0 if 'top' in attr or 'bottom' in attr else 2.5))

# Pas de numéro de page sur la couverture
sec0.different_first_page_header_footer = True
# Footer pages intérieures
fp = sec0.footer.paragraphs[0]
fp.clear()
rl = fp.add_run("Cyclistic Bike-Share  |  Sidi Mohamed ALLY  |  Google Data Analytics Certificate")
rl.font.size = Pt(8); rl.font.color.rgb = GRAY
fp.add_run("   ")
add_page_num(fp)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════
# Bandeau bleu (sans emoji pour compatibilité universelle)
band = doc.add_table(rows=1, cols=1)
band.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = band.cell(0, 0)
set_bg(bc, '#1A73E8')
bp = bc.paragraphs[0]; bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
br = bp.add_run("CYCLISTIC BIKE-SHARE")
br.font.size = Pt(26); br.font.bold = True; br.font.color.rgb = WHITE
bs = bc.add_paragraph(); bs.alignment = WD_ALIGN_PARAGRAPH.CENTER
bsr = bs.add_run("Rapport d'Analyse  ·  Case Study Google Data Analytics")
bsr.font.size = Pt(11); bsr.font.color.rgb = RGBColor(0xC5, 0xD9, 0xFD)
doc.add_paragraph()

# Titre & question
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Rapport d'Analyse — Case Study")
r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = DARK

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(
    "« Comment les membres annuels et les casual riders\n"
    "utilisent-ils les vélos Cyclistic différemment ? »"
)
r2.font.size = Pt(12); r2.font.italic = True; r2.font.color.rgb = GRAY
doc.add_paragraph()

# KPIs
kpis = [("3 885 439","Trajets analysés"),("64,6 %","Part Members"),
        ("35,4 %","Part Casual"),("14,6 min","Durée moy. Member"),
        ("37,1 min","Durée moy. Casual")]
kt = doc.add_table(rows=2, cols=5)
kt.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, (val, lbl) in enumerate(kpis):
    bg = '#D2E3FC' if ci % 2 else '#E8F0FE'
    set_bg(kt.cell(0, ci), bg); set_bg(kt.cell(1, ci), bg)
    pv = kt.cell(0, ci).paragraphs[0]; pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = pv.add_run(val); rv.font.bold = True; rv.font.size = Pt(14); rv.font.color.rgb = BLUE
    pl = kt.cell(1, ci).paragraphs[0]; pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl2 = pl.add_run(lbl); rl2.font.size = Pt(8); rl2.font.color.rgb = GRAY
doc.add_paragraph()

# Infos document
infos = [
    ("Analyste",          "Sidi Mohamed ALLY"),
    ("Titre",             "Analyste de données Professionnel Certifié"),
    ("Certification",     "Google Data Analytics Professional Certificate — Google / Coursera (13 juin 2026)"),
    ("Destinataire",      "Lily Moreno, Director of Marketing  |  Cyclistic Executive Team"),
    ("Date du rapport",   "Juin 2026"),
    ("Période analysée",  "Janvier 2019 – Décembre 2020"),
    ("Outils",            "Python (pandas · matplotlib · seaborn · plotly)  ·  SQL  ·  Tableau"),
]
it = doc.add_table(rows=len(infos), cols=2)
it.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (lbl, val) in enumerate(infos):
    c0, c1 = it.cell(ri, 0), it.cell(ri, 1)
    if ri % 2 == 0: set_bg(c0, '#F8F9FA'); set_bg(c1, '#F8F9FA')
    c0.width = Cm(4.2)
    # Mettre le nom en évidence
    r0 = c0.paragraphs[0].add_run(lbl)
    r0.bold = True; r0.font.color.rgb = BLUE; r0.font.size = Pt(9)
    r1 = c1.paragraphs[0].add_run(val)
    r1.font.size = Pt(9)
    if lbl == "Analyste":
        r1.font.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = DARK
    if lbl == "Titre":
        r1.font.italic = True; r1.font.color.rgb = GRAY
doc.add_paragraph()

# Badge + lien Credly
badge_row = doc.add_table(rows=1, cols=2)
badge_row.alignment = WD_TABLE_ALIGNMENT.CENTER
bc_img  = badge_row.cell(0, 0)
bc_link = badge_row.cell(0, 1)
bc_img.width  = Cm(4)
bc_link.width = Cm(12)
if os.path.exists(BADGE):
    pb = bc_img.paragraphs[0]; pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pb.add_run().add_picture(BADGE, width=Cm(3.2))
pl_c = bc_link.paragraphs[0]; pl_c.alignment = WD_ALIGN_PARAGRAPH.LEFT
pl_c.paragraph_format.space_before = Pt(10)
rl_title = pl_c.add_run("Google Data Analytics Professional Certificate\n")
rl_title.font.bold = True; rl_title.font.size = Pt(10); rl_title.font.color.rgb = DARK
rl_date = pl_c.add_run("Délivré le 13 juin 2026  ·  Émis par Google & Coursera\n")
rl_date.font.size = Pt(9); rl_date.font.color.rgb = GRAY
rl_v = pl_c.add_run("Vérifier le badge : ")
rl_v.font.size = Pt(9); rl_v.font.color.rgb = GRAY
add_hyperlink(pl_c, CREDLY, CREDLY, size=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# RÉSUMÉ EXÉCUTIF
# ══════════════════════════════════════════════════════════════════════════════
ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = ep.add_run("RÉSUMÉ EXÉCUTIF")
er.font.size = Pt(16); er.font.bold = True; er.font.color.rgb = BLUE
bottom_border(ep); doc.add_paragraph()

body(doc,
    "Cyclistic, programme de bike-share de Chicago, dispose de deux profils d'utilisateurs "
    "aux comportements opposés : les membres annuels (64,6 % des trajets) utilisent le service "
    "pour leurs déplacements domicile-travail, tandis que les casual riders (35,4 %) ont un usage "
    "récréatif et saisonnier. L'objectif stratégique est de convertir une partie des 1,37 million "
    "de casual riders en membres annuels, jugés plus rentables par les analystes financiers.",
    space_after=10)

colored_box(doc, "CONSTATS CLÉS", '#E8F0FE', [
    (None, "Les casual riders roulent 2,5× plus longtemps (37,1 min vs 14,6 min) mais 2× moins souvent → usage loisir, non utilitaire."),
    (None, "Double pic horaire membres à 8h et 17h (rush hours) vs pic unique casual à 15–17h → comportements structurellement différents."),
    (None, "51,5 % des trajets casual se concentrent en été (juin–août) → fenêtre de conversion estivale prioritaire."),
    (None, "Les casual se concentrent sur des stations touristiques (Lakefront, Navy Pier, Millennium Park) → ciblage géographique précis possible."),
], BLUE)

colored_box(doc, "TOP 3 RECOMMANDATIONS", '#E6F4EA', [
    ("1", "Campagne géolocalisée « Week-end → Abonnement » autour des stations touristiques."),
    ("2", "Offre « Été → Membre » : 2 mois offerts en mai–juin aux casual ayant 3+ trajets en 30 jours."),
    ("3", "Calculateur d'économies « domicile-travail » dans l'app pour démontrer la valeur de l'abonnement."),
], GREEN)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tr = tp.add_run("TABLE DES MATIÈRES")
tr.font.size = Pt(14); tr.font.bold = True; tr.font.color.rgb = BLUE
bottom_border(tp); doc.add_paragraph()

toc_entries = [
    ("1.", "Contexte & Problématique Business", False),
    ("2.", "Sources de données & Méthodologie", False),
    ("3.", "Nettoyage des données", False),
    ("4.", "Analyse & Insights", False),
    ("   4.1", "Vue d'ensemble : volume et durée", True),
    ("   4.2", "Comportement temporel", True),
    ("   4.3", "Types de vélos", True),
    ("   4.4", "Stations de départ — Casual riders", True),
    ("5.", "Recommandations Marketing", False),
    ("6.", "Conclusion & Limites", False),
    ("—",  "Signature de l'auteur", False),
    ("—",  "Glossaire", False),
    ("A.", "Annexe — Requêtes SQL", False),
    ("B.", "Annexe — Scripts Python clés", False),
    ("C.", "Annexe — Sources de données & Licence", False),
]
for num, title, sub in toc_entries:
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_after = Pt(2)
    p_t.paragraph_format.left_indent = Cm(0.5) if sub else Cm(0)
    rn = p_t.add_run(f"{num}  ")
    rn.font.bold = not sub; rn.font.size = Pt(10)
    rn.font.color.rgb = GRAY if sub else BLUE
    rt = p_t.add_run(title)
    rt.font.size = Pt(10); rt.font.color.rgb = GRAY if sub else DARK
doc.add_paragraph()

# ── Liste des figures ──────────────────────────────────────────────────────────
lf_p = doc.add_paragraph()
lf_r = lf_p.add_run("LISTE DES FIGURES")
lf_r.font.size = Pt(12); lf_r.font.bold = True; lf_r.font.color.rgb = BLUE
bottom_border(lf_p, 'E8EAED')
table(doc,
    ["N°", "Titre", "Fichier", "Section"],
    [
        ["Fig. 1", "Nombre total de trajets (membres vs casual)",       "01_rides_count.png",    "4.1"],
        ["Fig. 2", "Durée moyenne des trajets (minutes)",               "02_avg_ride_length.png","4.1"],
        ["Fig. 3", "Trajets par jour de la semaine",                    "03_rides_by_dow.png",   "4.2"],
        ["Fig. 4", "Évolution mensuelle des trajets (2019–2020)",       "04_rides_by_month.png", "4.2"],
        ["Fig. 5", "Répartition des trajets par heure de départ",       "07_rides_by_hour.png",  "4.2"],
        ["Fig. 6", "Durée moyenne par jour de la semaine",              "08_avg_length_by_dow.png","4.2"],
        ["Fig. 7", "Types de vélos utilisés (% par utilisateur)",       "05_bike_types.png",     "4.3"],
        ["Fig. 8", "Top 10 stations de départ — Casual riders",         "06_top10_stations.png", "4.4"],
    ],
    widths=[1.5, 7, 4.5, 2]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1.", "Contexte & Problématique Business")
body(doc,
    "Cyclistic est un programme de bike-share basé à Chicago comptant plus de 5 800 vélos "
    "géolocalisés répartis sur 692 stations. Lancé en 2016, le service propose une flexibilité "
    "tarifaire — pass journée, pass trajet unique ou abonnement annuel — qui a permis d'attirer "
    "un large public d'utilisateurs aux profils variés.")
body(doc,
    "Les analystes financiers ont établi que les membres annuels génèrent une rentabilité "
    "significativement supérieure. La Director of Marketing Lily Moreno a identifié une opportunité "
    "stratégique : maximiser la conversion des casual riders existants en membres annuels plutôt "
    "que de cibler exclusivement de nouveaux clients.")

h2(doc, "1.1", "Question d'analyse principale")
qt = doc.add_table(rows=1, cols=1); qt.alignment = WD_TABLE_ALIGNMENT.CENTER
qc = qt.cell(0, 0); set_bg(qc, '#F8F9FA')
qcp = qc.paragraphs[0]; qcp.alignment = WD_ALIGN_PARAGRAPH.CENTER
qcr = qcp.add_run(
    "« Comment les membres annuels et les casual riders utilisent-ils "
    "les vélos Cyclistic différemment ? »")
qcr.font.size = Pt(12); qcr.font.italic = True; qcr.font.color.rgb = BLUE
doc.add_paragraph()

h2(doc, "1.2", "Parties prenantes")
table(doc,
    ["Partie prenante", "Rôle", "Attente"],
    [["Lily Moreno", "Director of Marketing — commanditaire", "Recommandations actionnables basées sur les données"],
     ["Cyclistic Executive Team", "Décisionnaires finaux", "Validation via visualisations + données solides"],
     ["Marketing Analytics Team", "Implémentation", "Données exploitables pour les campagnes"]],
    widths=[4.5, 6, 5.5])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2.", "Sources de données & Méthodologie")
h2(doc, "2.1", "Fichiers utilisés")
table(doc,
    ["Fichier", "Période", "Lignes brutes"],
    [["Divvy_Trips_2019_Q1.csv", "Janv.–Mars 2019", "365 069"],
     ["Divvy_Trips_2020_Q1.csv", "Janv.–Mars 2020", "426 887"],
     ["202004-divvy-tripdata.csv", "Avril 2020", "84 776"],
     ["202005-divvy-tripdata.csv", "Mai 2020", "200 274"],
     ["202006-divvy-tripdata.csv", "Juin 2020", "343 005"],
     ["202007-divvy-tripdata.csv", "Juillet 2020", "551 480"],
     ["202008-divvy-tripdata.csv", "Août 2020", "622 361"],
     ["202009-divvy-tripdata.csv", "Septembre 2020", "532 958"],
     ["202010-divvy-tripdata.csv", "Octobre 2020", "388 653"],
     ["202011-divvy-tripdata.csv", "Novembre 2020", "259 716"],
     ["202012-divvy-tripdata.csv", "Décembre 2020", "131 573"],
     ["TOTAL", "Jan 2019 – Déc 2020", "3 906 752"]],
    widths=[7.5, 4, 4.5])

h2(doc, "2.2", "Crédibilité des données — ROCCC")
table(doc,
    ["Critère", "Statut", "Détail"],
    [["Reliable", "✅", "Données opérationnelles collectées automatiquement par les capteurs"],
     ["Original", "✅", "Source primaire : Motivate International Inc. — opérateur réel de Divvy"],
     ["Comprehensive", "✅", "Tous les trajets sur la période, aucun échantillonnage"],
     ["Current", "⚠️", "2019–2020 : tendances solides, mise à jour 2022+ recommandée avant déploiement"],
     ["Cited", "✅", "Divvy Data License Agreement — usage public autorisé"]],
    widths=[3.5, 1.5, 11])

h2(doc, "2.3", "Méthodologie — 6 phases Google Data Analytics")
phases = [
    ("Ask",     "Définir la question business et identifier les parties prenantes"),
    ("Prepare", "Collecter et documenter les 11 fichiers sources (data/01_data_description.md)"),
    ("Process", "Nettoyage, harmonisation et consolidation (scripts/02_cleaning.py)"),
    ("Analyze", "Analyse descriptive Python + 12 requêtes SQL (03_analysis.py, 03_queries.sql)"),
    ("Share",   "8 graphiques statiques (matplotlib/seaborn) + dashboard Plotly interactif"),
    ("Act",     "Top 3 recommandations + rapport complet + export Tableau"),
]
pt = doc.add_table(rows=len(phases), cols=2)
pt.style = 'Table Grid'; pt.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (ph, desc) in enumerate(phases):
    c0, c1 = pt.cell(ri, 0), pt.cell(ri, 1)
    set_bg(c0, '#E8F0FE'); c0.width = Cm(2.8)
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(ph); r0.bold = True; r0.font.color.rgb = BLUE; r0.font.size = Pt(10)
    r1 = c1.paragraphs[0].add_run(desc); r1.font.size = Pt(9)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3.", "Nettoyage des données")
body(doc, "Réalisé intégralement en Python (pandas) via scripts/02_cleaning.py. "
          "Log détaillé disponible dans data/02_cleaning_log.md.")

h2(doc, "3.1", "Transformations appliquées")
table(doc,
    ["Étape", "Action", "Impact"],
    [["Harmonisation schéma",    "Renommage colonnes 2019 Q1 + Subscriber→member / Customer→casual", "Unification des 11 fichiers"],
     ["Conversion datetime",     "started_at et ended_at → datetime64",                               "Calculs temporels précis"],
     ["ride_length",             "(ended_at − started_at) en minutes",                                 "Variable d'analyse principale"],
     ["Variables temporelles",   "day_of_week, month, year, season, hour_of_day",                     "5 nouvelles colonnes"],
     ["Doublons",                "Suppression sur ride_id",                                            "Dédoublonnage"],
     ["ride_length ≤ 0 min",     "Suppression (erreurs système)",                                      "Données invalides écartées"],
     ["ride_length > 1 440 min", "Suppression (vélos non retournés > 24h)",                            "Outliers extrêmes écartés"],
     ["Stations maintenance",    "Suppression : HQ QR, TEST, DIVVY, Hubbard Bike-checking",           "Trajets internes exclus"],
     ["Résultat final",          "3 906 752 → 3 885 439 trajets propres",                             "−21 313 lignes (−0,5 %)"]],
    widths=[4, 8, 4])

h2(doc, "3.2", "Contraintes & limites")
for lim in [
    "Confidentialité : pas d'identifiant personnel (PII) — impossible de tracer les trajets d'un même utilisateur.",
    "Données manquantes : 5–15 % de valeurs nulles sur les noms de stations (vélos electric/docked).",
    "Schéma hétérogène : le fichier 2019 Q1 utilise une structure différente des fichiers 2020.",
    "Impact COVID-19 : creux d'avril 2020 atypique lié au confinement.",
    "Données historiques : tendances solides mais une mise à jour 2022–2024 serait recommandée.",
]:
    bullet(doc, lim)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4.", "Analyse & Insights")

# 4.1
h2(doc, "4.1", "Vue d'ensemble : volume et durée")
body(doc,
    "Sur 3 885 439 trajets analysés, les membres représentent 64,6 % du volume mais effectuent "
    "des déplacements 2,5× plus courts que les casual riders. Cette asymétrie révèle des usages "
    "fondamentalement distincts.")
table(doc,
    ["Métrique", "Members", "Casual riders", "Ratio"],
    [["Nombre de trajets", "2 508 757", "1 376 682", "Members × 1,8"],
     ["Part du total", "64,6 %", "35,4 %", "—"],
     ["Durée moyenne", "14,6 min", "37,1 min", "Casual × 2,5"],
     ["Durée médiane", "10,6 min", "21,7 min", "Casual × 2,0"],
     ["Jour de pointe", "Jeudi", "Samedi", "—"],
     ["Usage principal", "Domicile ↔ Travail", "Loisirs / Tourisme", "—"]],
    widths=[5, 3.5, 3.5, 4])

ft = doc.add_table(rows=1, cols=2); ft.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, (fname, num, cap) in enumerate([
    ("01_rides_count.png",     "1", "Nombre total de trajets"),
    ("02_avg_ride_length.png", "2", "Durée moyenne (minutes)"),
]):
    cell = ft.cell(0, ci)
    path = os.path.join(FIG_DIR, fname)
    if os.path.exists(path):
        pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.add_run().add_picture(path, width=Cm(7.5))
    pc = cell.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run(f"Fig. {num} — {cap}  ({fname})")
    rc.font.size = Pt(8); rc.font.italic = True; rc.font.color.rgb = GRAY
doc.add_paragraph()
insight_box(doc,
    "Les casual riders effectuent 2,5× plus de minutes de vélo par trajet que les membres. "
    "Ce n'est pas un usage utilitaire — c'est un comportement récréatif ou touristique structurel.")

doc.add_page_break()

# 4.2
h2(doc, "4.2", "Comportement temporel")
body(doc,
    "L'analyse temporelle révèle la différence structurelle la plus significative : les membres ont "
    "un usage utilitaire rythmé par la semaine de travail, les casual riders un usage récréatif "
    "concentré sur le week-end et les heures de l'après-midi.")

body(doc, "Par jour de la semaine :", space_after=2)
bullet(doc, "Members : pic le jeudi (~391k trajets), volume stable lundi–vendredi → usage domicile-travail.")
bullet(doc, "Casual : pic le samedi (~318k trajets), profil en cloche autour du week-end → usage loisir.")
figure(doc, "03_rides_by_dow.png", "3", "Trajets par jour de la semaine (membres vs casual)")

body(doc, "Par mois et saison :", space_after=2)
bullet(doc, "Pic estival commun (juin–août), bien plus prononcé chez les casual.")
bullet(doc, "Casual hyper-saisonniers : 51,5 % de leurs trajets annuels en été vs 32 % pour les membres.")
bullet(doc, "Hiver : membres actifs (556k) vs casual quasi-absents (57k).")
bullet(doc, "Creux d'avril 2020 visible pour les deux types (confinement COVID-19).")
figure(doc, "04_rides_by_month.png", "4", "Évolution mensuelle des trajets (Janv. 2019 – Déc. 2020)", width=15)

doc.add_page_break()

body(doc, "Par heure de départ :", space_after=2)
bullet(doc, "Members : double pic à 8h et 17h → trajets domicile-travail confirmés.")
bullet(doc, "Casual : montée progressive dès 10h, pic unique à 15–17h → loisir, pas d'urgence horaire.")
figure(doc, "07_rides_by_hour.png", "5", "Répartition des trajets par heure de départ")

body(doc, "Durée par jour de la semaine :", space_after=2)
body(doc,
    "Les casual maintiennent des trajets longs toute la semaine (>25 min même en semaine), "
    "ce qui exclut une explication liée uniquement au week-end.")
figure(doc, "08_avg_length_by_dow.png", "6", "Durée moyenne des trajets par jour de la semaine")

insight_box(doc,
    "Trois signaux temporels convergent — week-end, après-midi (15–17h), été (juin–août). "
    "Ce sont les trois fenêtres d'intervention marketing prioritaires pour cibler les casual riders.")

# 4.3
h2(doc, "4.3", "Types de vélos utilisés")
body(doc,
    "Les docked bikes dominent pour les deux types (>83 %). Les casual riders utilisent "
    "proportionnellement plus d'electric bikes (15,2 % vs 11,8 % pour les membres), ce qui peut "
    "indiquer une sensibilité à l'effort ou un usage sur de plus longues distances touristiques.")
figure(doc, "05_bike_types.png", "7", "Types de vélos utilisés — % par catégorie d'utilisateur")
insight_box(doc,
    "La préférence casual pour l'electric bike (15,2 %) ouvre une piste de communication "
    "sur le confort et la polyvalence — particulièrement pertinente pour les navetteurs potentiels.")

# 4.4
h2(doc, "4.4", "Stations de départ — Casual riders")
body(doc,
    "Les 10 stations les plus fréquentées par les casual riders sont toutes situées dans des "
    "zones touristiques et de loisirs : Lakefront Trail, Navy Pier, Millennium Park, Shedd Aquarium. "
    "Cette concentration géographique rend un ciblage marketing géolocalisé très précis possible.")
figure(doc, "06_top10_stations.png", "8", "Top 10 stations de départ — Casual riders (cibles prioritaires)")
insight_box(doc,
    "Les 10 stations prioritaires sont toutes concentrées sur le Lakefront de Chicago "
    "— une zone de ciblage géographique très précise, idéale pour des campagnes géolocalisées.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5.", "Recommandations Marketing")
body(doc,
    "Les trois recommandations suivantes découlent directement des insights de l'analyse. "
    "Elles sont ordonnées par ratio Impact/Effort décroissant.")

colored_box(doc, "RECOMMANDATION 1 — Campagne « Week-end → Abonnement » géolocalisée", '#E8F0FE', [
    ("Insight",        "Les casual riders se concentrent le week-end sur des stations touristiques identifiables."),
    ("Action",         "Publicités géolocalisées dans un rayon de 500 m autour des 10 stations casual prioritaires, vendredi soir et samedi matin."),
    ("Message clé",    "« Tu roules déjà le week-end — économise avec l'abonnement annuel. »"),
    ("Canaux",         "Instagram/TikTok géolocalisés · notifications push in-app · affichage en station."),
    ("Impact attendu", "Conversion des casual récurrents du week-end au moment de leur usage maximal."),
], BLUE)

colored_box(doc, "RECOMMANDATION 2 — Offre « Été → Membre » à tarif préférentiel", '#E6F4EA', [
    ("Insight",        "51,5 % des trajets casual en été → fenêtre de conversion optimale."),
    ("Action",         "En mai–juin : 2 premiers mois offerts aux casual ayant effectué 3+ trajets en 30 jours. Email personnalisé avec calcul d'économies réalisées."),
    ("Message clé",    "« Tu as utilisé Cyclistic X fois ce mois-ci — tu aurais économisé Y€ avec un abonnement. »"),
    ("Canaux",         "Email ciblé · notification push in-app · bannières dans l'app."),
    ("Impact attendu", "Conversion pendant le pic d'usage, avant l'inactivité automne/hiver."),
], GREEN)

colored_box(doc, "RECOMMANDATION 3 — Calculateur d'économies « domicile-travail »", '#FEF7E0', [
    ("Insight",        "Les membres utilisent les vélos aux heures de rush (8h, 17h). Les casual ignorent peut-être cet usage."),
    ("Action",         "Calculateur interactif dans l'app : « Si tu fais X trajets/semaine, l'abonnement te coûte Y€/trajet vs Z€ en pass journée. »"),
    ("Canaux",         "Content marketing · affichage stations aux heures de rush."),
    ("Impact attendu", "Démonstration concrète de la valeur économique de l'abonnement."),
], AMBER)

h2(doc, "5.4", "Matrice de priorisation — Effort × Impact")
table(doc,
    ["Recommandation", "Effort", "Impact estimé", "Délai", "KPI de mesure"],
    [["1 — Campagne géolocalisée", "Faible", "Élevé ★★★★", "1–2 semaines",  "Taux de conversion par station"],
     ["2 — Offre été Membre",      "Moyen",  "Très élevé ★★★★★", "2–4 semaines", "Nouveaux membres juin–août"],
     ["3 — Calculateur économies", "Moyen",  "Moyen ★★★", "4–8 semaines",  "CTR calculateur & conversions"]],
    widths=[4.5, 2, 3, 2.5, 4])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "6.", "Conclusion & Limites")
h2(doc, "6.1", "Synthèse")
body(doc,
    "Cette analyse de 3 885 439 trajets (2019–2020) établit une distinction comportementale nette "
    "et cohérente entre les deux types d'utilisateurs Cyclistic. Les membres ont un usage utilitaire "
    "et régulier, tandis que les casual riders ont un usage récréatif, saisonnier et géographiquement "
    "concentré. La convergence de trois signaux temporels (week-end, après-midi, été) et d'un signal "
    "géographique (Lakefront) offre un cadre de ciblage marketing très précis.")
body(doc,
    "Les casual riders les plus réguliers sont des candidats naturels à la conversion, à condition "
    "que la communication leur démontre la valeur de l'abonnement dans leur contexte d'usage réel.")

h2(doc, "6.2", "Limites & prochaines étapes")
for lim in [
    "Absence de données individuelles : impossible d'identifier les casual récurrents sans accès aux comptes utilisateurs.",
    "Période historique : une mise à jour avec des données 2022–2024 est recommandée avant le déploiement des campagnes.",
    "Absence de données démographiques : âge, genre et localisation résidentielle permettraient un ciblage plus fin.",
    "Recommandation opérationnelle : A/B test sur un sous-groupe de casual riders (1 offre, 1 canal, 1 station pilote) avant déploiement à grande échelle.",
]:
    bullet(doc, lim)

doc.add_paragraph()
fq = doc.add_paragraph(); fq.alignment = WD_ALIGN_PARAGRAPH.CENTER
fq.paragraph_format.space_before = Pt(12)
rfq = fq.add_run(
    "« Les données ne prennent de valeur que lorsqu'elles conduisent à des décisions. "
    "Ce rapport a été conçu pour être directement actionnable. »")
rfq.font.italic = True; rfq.font.size = Pt(10); rfq.font.color.rgb = GRAY
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SIGNATURE
# ══════════════════════════════════════════════════════════════════════════════
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("RÉALISÉ PAR")
sr.font.size = Pt(11); sr.font.bold = True; sr.font.color.rgb = BLUE
bottom_border(sp)
doc.add_paragraph()

sig_tbl = doc.add_table(rows=1, cols=2)
sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
sc_badge = sig_tbl.cell(0, 0); sc_info = sig_tbl.cell(0, 1)
sc_badge.width = Cm(4.5); sc_info.width = Cm(11.5)

if os.path.exists(BADGE):
    pb2 = sc_badge.paragraphs[0]; pb2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pb2.add_run().add_picture(BADGE, width=Cm(3.8))

pi = sc_info.paragraphs[0]
pi.paragraph_format.space_before = Pt(8)
rname = pi.add_run("Sidi Mohamed ALLY\n")
rname.font.bold = True; rname.font.size = Pt(16); rname.font.color.rgb = DARK
rtitle = pi.add_run("Analyste de données Professionnel Certifié\n")
rtitle.font.size = Pt(11); rtitle.font.italic = True; rtitle.font.color.rgb = GRAY
rcert = pi.add_run("Google Data Analytics Professional Certificate\n")
rcert.font.size = Pt(10); rcert.font.color.rgb = BLUE; rcert.font.bold = True
rdate = pi.add_run("Certifié le 13 juin 2026  ·  Délivré par Google & Coursera\n")
rdate.font.size = Pt(9); rdate.font.color.rgb = GRAY
pi.add_run("Badge de certification vérifiable : ")
add_hyperlink(pi, CREDLY, CREDLY, size=9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# GLOSSAIRE
# ══════════════════════════════════════════════════════════════════════════════
gp = doc.add_paragraph()
gr = gp.add_run("GLOSSAIRE")
gr.font.size = Pt(14); gr.font.bold = True; gr.font.color.rgb = BLUE
bottom_border(gp)
doc.add_paragraph()

glossaire = [
    ("A/B Test",       "Expérimentation marketing comparant deux variantes (A et B) pour mesurer laquelle est la plus efficace."),
    ("Casual rider",   "Utilisateur Cyclistic ayant souscrit un pass journée ou un pass trajet unique (non-membre)."),
    ("Credly",         "Plateforme de badges de certification numériques vérifiables en ligne."),
    ("CTR",            "Click-Through Rate — taux de clic sur une publicité ou un élément interactif."),
    ("CSV",            "Comma-Separated Values — format de fichier texte pour les données tabulaires."),
    ("KPI",            "Key Performance Indicator — indicateur clé de performance permettant de mesurer l'atteinte d'un objectif."),
    ("Member",         "Utilisateur Cyclistic détenteur d'un abonnement annuel."),
    ("PII",            "Personally Identifiable Information — toute donnée permettant d'identifier une personne physique."),
    ("ROCCC",          "Cadre d'évaluation de la qualité des données : Reliable, Original, Comprehensive, Current, Cited."),
    ("Rush hours",     "Heures de pointe des transports (7h–9h et 16h–18h) caractérisées par un trafic intense."),
    ("SQL",            "Structured Query Language — langage de requête pour les bases de données relationnelles."),
]
table(doc,
    ["Terme", "Définition"],
    [[t, d] for t, d in glossaire],
    widths=[3.5, 12.5])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ANNEXE A — SQL
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "A.", "Annexe — Requêtes SQL")
body(doc,
    "Les 12 requêtes SQL ci-dessous reproduisent l'analyse en SQL standard (PostgreSQL / BigQuery / DuckDB). "
    "Note de compatibilité : PERCENTILE_CONT est utilisé à la place de MEDIAN() "
    "pour assurer la portabilité sur les moteurs SQL majeurs.")

with open(SQL_FILE, encoding='utf-8') as f:
    sql_raw = f.read()

for block in re.split(r'\n(?=-- ──)', sql_raw):
    block = block.strip()
    if not block:
        continue
    lines_b = block.split('\n')
    title_l = next((l for l in lines_b if '──' in l), None)
    if title_l:
        pt2 = doc.add_paragraph()
        pt2.paragraph_format.space_before = Pt(8)
        rt2 = pt2.add_run(title_l.replace('--', '').replace('─', '').strip())
        rt2.font.bold = True; rt2.font.size = Pt(9); rt2.font.color.rgb = BLUE
    pc2 = doc.add_paragraph()
    pc2.paragraph_format.left_indent = Cm(0.5)
    pc2.paragraph_format.space_after = Pt(4)
    rc2 = pc2.add_run(block)
    rc2.font.name = 'Courier New'; rc2.font.size = Pt(7.5); rc2.font.color.rgb = DARK
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ANNEXE B — SCRIPTS PYTHON CLÉS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "B.", "Annexe — Scripts Python clés")
body(doc,
    "Extraits des scripts de nettoyage et d'analyse illustrant les principales décisions "
    "techniques. Scripts complets disponibles dans scripts/.")

# Extrait 1 : load_legacy_2019
code_block(doc, """\
# scripts/02_cleaning.py — Mapping du schéma Legacy 2019
COLS_2020 = ["ride_id", "rideable_type", "started_at", "ended_at",
             "start_station_name", "start_station_id",
             "end_station_name", "end_station_id",
             "start_lat", "start_lng", "end_lat", "end_lng", "member_casual"]

def load_legacy_2019(path: str) -> pd.DataFrame:
    df = pd.read_csv(path, low_memory=False)
    df = df.rename(columns={
        "trip_id": "ride_id", "start_time": "started_at", "end_time": "ended_at",
        "from_station_id": "start_station_id", "from_station_name": "start_station_name",
        "to_station_id": "end_station_id", "to_station_name": "end_station_name",
        "usertype": "member_casual",
    })
    df["member_casual"] = df["member_casual"].replace({"Subscriber": "member", "Customer": "casual"})
    df["rideable_type"] = "unknown"
    df[["start_lat","start_lng","end_lat","end_lng"]] = np.nan
    return df[COLS_2020]""",
    title="B.1 — Harmonisation du schéma Legacy 2019 (02_cleaning.py)")

# Extrait 2 : pipeline nettoyage
code_block(doc, """\
# scripts/02_cleaning.py — Pipeline de nettoyage
df["started_at"]  = pd.to_datetime(df["started_at"], infer_datetime_format=True, errors="coerce")
df["ended_at"]    = pd.to_datetime(df["ended_at"],   infer_datetime_format=True, errors="coerce")
df["ride_length"] = (df["ended_at"] - df["started_at"]).dt.total_seconds() / 60
df["day_of_week"] = df["started_at"].dt.dayofweek.map({6:1,0:2,1:3,2:4,3:5,4:6,5:7})
df["month"]       = df["started_at"].dt.month
df["year"]        = df["started_at"].dt.year
df["hour_of_day"] = df["started_at"].dt.hour
df["season"]      = df["month"].map({12:"Winter",1:"Winter",2:"Winter",
                                      3:"Spring",4:"Spring",5:"Spring",
                                      6:"Summer",7:"Summer",8:"Summer",
                                      9:"Fall",10:"Fall",11:"Fall"})
df = df.drop_duplicates(subset=["ride_id"])
df = df[df["ride_length"] > 0]
df = df[df["ride_length"] <= 1440]
mask_hq = (df["start_station_name"].str.contains("HQ QR|DIVVY|TEST", case=False, na=False) |
           df["end_station_name"].str.contains("HQ QR|DIVVY|TEST", case=False, na=False))
df = df[~mask_hq]""",
    title="B.2 — Pipeline de nettoyage (02_cleaning.py)")

# Extrait 3 : analyse
code_block(doc, """\
# scripts/03_analysis.py — Statistiques descriptives & analyse clé
desc = df.groupby("member_casual")["ride_length"].agg(
    mean_ride_length="mean", median_ride_length="median",
    max_ride_length="max", total_rides="count"
).round(2).reset_index()

# Top 10 stations casual (cibles marketing)
top10_casual = (df[(df["start_station_name"].notna()) & (df["member_casual"]=="casual")]
                .groupby("start_station_name").size().reset_index(name="total_rides")
                .sort_values("total_rides", ascending=False)
                .head(10).reset_index(drop=True))

# Pourcentage saisonnier
seasonal = df.groupby(["member_casual","season"]).size().reset_index(name="total_rides")
seasonal["pct"] = (seasonal["total_rides"] /
                   seasonal.groupby("member_casual")["total_rides"].transform("sum") * 100).round(1)""",
    title="B.3 — Analyse descriptive & top 10 stations (03_analysis.py)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ANNEXE C — SOURCES & LICENCE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "C.", "Annexe — Sources de données & Licence")
table(doc,
    ["Élément", "Détail"],
    [["Source",   "Divvy Trip Data — Motivate International Inc."],
     ["Licence",  "Divvy Data License Agreement — usage public autorisé"],
     ["Période",  "Janvier 2019 – Décembre 2020 (11 fichiers CSV)"],
     ["Volume",   "3 906 752 trajets bruts → 3 885 439 après nettoyage"],
     ["Anonymat", "Aucune PII — données 100 % anonymisées"],
     ["Note",     "Cyclistic est une entreprise fictive (Capstone Case Study 1, Google DA Certificate). "
                  "Les données Divvy sont réelles et adaptées à des fins pédagogiques."]],
    widths=[3.5, 12.5])

body(doc, "Livrables du projet :", space_after=3)
for item in [
    "scripts/02_cleaning.py      — Nettoyage & consolidation (pandas)",
    "scripts/03_analysis.py      — Analyse descriptive Python",
    "scripts/03_queries.sql      — 12 requêtes SQL (PostgreSQL / BigQuery / DuckDB)",
    "scripts/04_visualizations.py — 8 graphiques PNG (matplotlib / seaborn)",
    "scripts/05_dashboard.py     — Dashboard interactif Plotly",
    "scripts/06_tableau_export.py — Export CSV pour Tableau Public",
    "05_dashboard.html           — Dashboard interactif (ouvrir dans un navigateur)",
    "06_report.pdf               — Ce rapport en version PDF",
]:
    p_i = doc.add_paragraph()
    p_i.paragraph_format.left_indent = Cm(0.5)
    p_i.paragraph_format.space_after = Pt(2)
    ri2 = p_i.add_run(item)
    ri2.font.name = 'Courier New'; ri2.font.size = Pt(8.5)

# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT_DOCX)
size_kb = os.path.getsize(OUT_DOCX) // 1024
print(f"DOCX genere : {OUT_DOCX}")
print(f"Taille      : {size_kb} KB")
